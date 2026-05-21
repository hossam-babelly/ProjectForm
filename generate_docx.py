#!/usr/bin/env python3
"""
Complete Word document generator with SmartArt org chart.
Called from Node.js: python3 generate_docx.py <data.json> <output.docx>
"""
import sys, json, os, uuid, zipfile, shutil, re
from copy import deepcopy

# python-docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── Color helpers ──────────────────────────────────────────────
def hex2rgb(h): h=h.lstrip('#'); return tuple(int(h[i:i+2],16) for i in (0,2,4))
def rgb(h): r,g,b=hex2rgb(h); return RGBColor(r,g,b)

C = {
    'DARK':    '1F3864',
    'GREY':    'AEAAAA',
    'HDR_BLK': '262626',
    'HDR_BLU': 'D9E2F3',
    'P1D':     '8EA9DB', 'P1O': 'D9E1F2', 'P1E': 'B4C6E7',
    'P2D':     'F4B084', 'P2O': 'FCE4D6', 'P2E': 'F8CBAD',
    'P3D':     '8EAADB',
    'SUB':     '808080',
    'SLBL':    'DEEAF6', 'SVAL': 'FBE4D5',
    'RODD':    'FBE4D5', 'REVEN': 'F2F2F2',
    'WHITE':   'FFFFFF', 'BLACK': '000000',
    'ORG_ADM': '4472C4', 'ORG_EXE': '70AD47',
    'ORG_SVC': 'ED7D31', 'ORG_OTH': '5B9BD5',
}

PROD_PAL = [
    {'d': C['P1D'], 'o': C['P1O'], 'e': C['P1E']},
    {'d': C['P2D'], 'o': C['P2O'], 'e': C['P2E']},
    {'d': C['P3D'], 'o': C['P1O'], 'e': C['P1E']},
]

MONTHS = ['شهر 1','شهر 2','شهر 3','شهر 4','شهر 5','شهر 6',
          'شهر 7','شهر 8','شهر 9','شهر 10','شهر 11','شهر 12']

FONT_NAME = 'Sakkal Majalla'
FONT_SIZE = 14  # pt

def fM(v):
    try:
        n = float(str(v).replace('$','').replace(',','').strip())
        if n == 0: return '0 $'
        return f"{n:,.0f} $" if n == int(n) else f"{n:,.2f} $"
    except: return '0 $'

def fN(v):
    try:
        n = float(str(v).replace('$','').replace(',','').strip())
        if n == 0: return '0'
        return f"{n:,.0f}" if n == int(n) else f"{n:,.2f}"
    except: return '0'

# ── XML cell shading helper ────────────────────────────────────
def set_cell_shading(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color.upper())
    # Remove existing shd
    for old in tcPr.findall(qn('w:shd')): tcPr.remove(old)
    tcPr.append(shd)

def set_cell_borders(cell, color='000000', sz=8):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')): tcPr.remove(old)
    bdr = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        bdr.append(el)
    tcPr.append(bdr)

def set_cell_rtl(cell):
    """Force RTL reading order on cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # textDirection
    td = OxmlElement('w:textDirection')
    td.set(qn('w:val'), 'btLr')
    # Actually for RTL we need rtl on paragraph
    pass

def set_para_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    # remove existing
    for old in pPr.findall(qn('w:bidi')): pPr.remove(old)
    pPr.append(bidi)
    jc = pPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        pPr.append(jc)
    jc.set(qn('w:val'), 'center')

def add_cell(cell, text, fill, text_color='FFFFFF', bold=True, align='center'):
    set_cell_shading(cell, fill)
    set_cell_borders(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Clear existing
    for p in cell.paragraphs: p._p.getparent().remove(p._p)
    p = cell.add_paragraph()
    set_para_rtl(p)
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text) if text is not None else '')
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE)
    run.font.color.rgb = rgb(text_color)
    # Set RTL on run
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    for old in rPr.findall(qn('w:rtl')): rPr.remove(old)
    rPr.append(rtl)
    return p

def set_table_bidi(table):
    """Set bidiVisual on table"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:bidiVisual')): tblPr.remove(old)
    bidi = OxmlElement('w:bidiVisual')
    tblPr.append(bidi)

def set_col_widths(table, widths_cm):
    """Set column widths"""
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
                tcW = OxmlElement('w:tcW')
                # Convert cm to twips (1 cm = 567 twips)
                twips = int(widths_cm[j] * 567)
                tcW.set(qn('w:w'), str(twips))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.insert(0, tcW)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_breaks.WD_BREAK.PAGE)
    return p

# ── Section header row ──────────────────────────────────────────
def section_hdr_row(table, title, fill=None, text_color='FFFFFF'):
    if fill is None: fill = C['DARK']
    row = table.add_row()
    # Merge all cells
    cells = row.cells
    if len(cells) > 1:
        cells[0].merge(cells[-1])
    add_cell(cells[0], title, fill, text_color, bold=True, align='center')
    return row

def col_hdr_row(table, labels, fill=None, text_color='FFFFFF'):
    if fill is None: fill = C['HDR_BLK']
    row = table.add_row()
    for i, lbl in enumerate(labels):
        if i < len(row.cells):
            add_cell(row.cells[i], lbl, fill, text_color, bold=True)
    return row

def data_row(table, values, fill=None, text_color='000000', bold=True):
    if fill is None: fill = C['WHITE']
    row = table.add_row()
    for i, val in enumerate(values):
        if i < len(row.cells):
            add_cell(row.cells[i], val, fill, text_color, bold=bold)
    return row

def total_row(table, label, value, n_cols, fill=None):
    if fill is None: fill = C['DARK']
    row = table.add_row()
    cells = row.cells
    # Merge all except last
    if n_cols > 2:
        cells[0].merge(cells[n_cols-2])
    add_cell(cells[0], label, fill, C['WHITE'], bold=True, align='center')
    add_cell(cells[-1], value, fill, C['WHITE'], bold=True, align='center')
    return row

# ════════════════════════════════════════════════════════════════
# SMARTART ORG CHART GENERATOR
# ════════════════════════════════════════════════════════════════

def generate_smartart_data_xml(hr_rows):
    """Generate SmartArt data1.xml with org hierarchy from hr_rows"""
    
    # Build employee list, expanding qty>1 into multiple nodes
    employees = []
    for r in hr_rows:
        pos = r.get('position','')
        typ = r.get('type','')
        qty = int(r.get('qty', 1) or 1)
        reports = r.get('reports','')
        if not pos: continue
        for _ in range(qty):
            employees.append({'pos': pos, 'type': typ, 'reports': reports})
    
    # Assign UUIDs
    for e in employees:
        e['id'] = '{' + str(uuid.uuid4()).upper() + '}'
    
    # Build XML
    ns = {
        'dgm': 'http://schemas.openxmlformats.org/drawingml/2006/diagram',
        'a':   'http://schemas.openxmlformats.org/drawingml/2006/main',
    }
    
    def uid(): return '{' + str(uuid.uuid4()).upper() + '}'
    
    lines = ['<?xml version="1.0" encoding="UTF-8" standalone="yes"?>']
    lines.append('<dgm:dataModel xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">')
    lines.append('<dgm:ptLst>')
    
    # Doc node
    doc_id = uid()
    lines.append(f'<dgm:pt modelId="{doc_id}" type="doc"><dgm:prSet/><dgm:spPr/><dgm:t><a:bodyPr/><a:lstStyle/><a:p><a:r><a:t/></a:r></a:p></dgm:t></dgm:pt>')
    
    # Employee nodes
    for e in employees:
        eid = e['id']
        pos = e['pos']
        typ = e['type']
        # Map type to fill color using theme
        bg_map = {'إداري':'4472C4','تنفيذي':'70AD47','مزود خدمة':'ED7D31'}
        bg = bg_map.get(typ, '5B9BD5')
        lines.append(f'''<dgm:pt modelId="{eid}" type="asst">
<dgm:prSet><dgm:presLayoutVars><dgm:var type="bulletEnabled" val="0"/></dgm:presLayoutVars></dgm:prSet>
<dgm:spPr><a:solidFill><a:srgbClr val="{bg}"/></a:solidFill></dgm:spPr>
<dgm:t><a:bodyPr/><a:lstStyle/><a:p><a:r><a:rPr lang="ar-SA" rtl="1"/><a:t>{pos}</a:t></a:r></a:p>
<a:p><a:r><a:rPr lang="ar-SA" rtl="1"/><a:t>({typ})</a:t></a:r></a:p></dgm:t></dgm:pt>''')
    
    lines.append('</dgm:ptLst>')
    lines.append('<dgm:cxnLst>')
    
    # Connections: doc -> root, then parent -> child
    # Find roots (no reports or reports to unknown)
    all_pos = [e['pos'] for e in employees]
    
    def find_emp_by_pos(pos):
        return [e for e in employees if e['pos'] == pos]
    
    for e in employees:
        reports = e.get('reports','')
        if not reports or reports not in all_pos:
            # Connect doc -> this employee (root)
            cid = uid()
            lines.append(f'<dgm:cxn modelId="{cid}" srcId="{doc_id}" destId="{e["id"]}" srcOrd="0" destOrd="0"/>')
        else:
            # Find a parent
            parents = find_emp_by_pos(reports)
            if parents:
                parent = parents[0]
                cid = uid()
                lines.append(f'<dgm:cxn modelId="{cid}" srcId="{parent["id"]}" destId="{e["id"]}" srcOrd="0" destOrd="0"/>')
    
    lines.append('</dgm:cxnLst>')
    lines.append('<dgm:bg/>')
    lines.append('<dgm:whole/>')
    lines.append('</dgm:dataModel>')
    
    return '\n'.join(lines)


def embed_smartart_in_docx(doc, hr_rows, output_path):
    """
    Save doc to output_path, then add SmartArt by manipulating the zip.
    SmartArt will use the layout/colors/style from the original target file.
    """
    import tempfile
    
    # Save current doc to temp
    tmp = tempfile.mktemp(suffix='.docx')
    doc.save(tmp)
    
    # Generate new data XML
    new_data_xml = generate_smartart_data_xml(hr_rows)
    
    # Read original SmartArt files from target
    smartart_files = {
        'layout1.xml':     open('/tmp/docx_target/word/diagrams/layout1.xml','rb').read(),
        'colors1.xml':     open('/tmp/docx_target/word/diagrams/colors1.xml','rb').read(),
        'quickStyle1.xml': open('/tmp/docx_target/word/diagrams/quickStyle1.xml','rb').read(),
        'data1.xml':       new_data_xml.encode('utf-8'),
    }
    
    # Read the SmartArt container XML (the paragraph with drawing)
    # We'll create a simple inline drawing that references the SmartArt
    smartart_para_xml = '''<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
         xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
         xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
         xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
         xmlns:dgm="http://schemas.openxmlformats.org/drawingml/2006/diagram"
         xmlns:dsp="http://schemas.microsoft.com/office/drawing/2008/diagram">
  <w:pPr><w:bidi/><w:jc w:val="center"/></w:pPr>
  <w:r>
    <w:rPr/>
    <w:drawing>
      <wp:inline distT="0" distB="0" distL="0" distR="0">
        <wp:extent cx="9144000" cy="5715000"/>
        <wp:effectExtent l="0" t="0" r="0" b="0"/>
        <wp:docPr id="1" name="SmartArt 1"/>
        <wp:cNvGraphicFramePr/>
        <a:graphic>
          <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/diagram">
            <dgm:relIds r:dm="rId_dm" r:lo="rId_lo" r:qs="rId_qs" r:cs="rId_cs" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>
          </a:graphicData>
        </a:graphic>
      </wp:inline>
    </w:drawing>
  </w:r>
</w:p>'''
    
    # Manipulate the docx zip
    import zipfile as zf
    
    with zf.ZipFile(tmp, 'r') as zin, zf.ZipFile(output_path, 'w', zf.ZIP_DEFLATED) as zout:
        # Copy all existing files
        for item in zin.infolist():
            zout.writestr(item, zin.read(item.filename))
        
        # Add SmartArt diagram files
        for fname, content in smartart_files.items():
            zout.writestr(f'word/diagrams/{fname}', content)
        
        # Add relationship for diagrams in document.xml.rels
        # Read existing rels
        try:
            rels_content = zin.read('word/_rels/document.xml.rels').decode('utf-8')
        except:
            rels_content = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
</Relationships>'''
        
        # Add diagram relationships
        rels_content = rels_content.replace('</Relationships>',
            '''<Relationship Id="rId_dm" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramData" Target="diagrams/data1.xml"/>
<Relationship Id="rId_lo" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramLayout" Target="diagrams/layout1.xml"/>
<Relationship Id="rId_qs" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramQuickStyle" Target="diagrams/quickStyle1.xml"/>
<Relationship Id="rId_cs" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramColors" Target="diagrams/colors1.xml"/>
</Relationships>''')
        
        # We need to overwrite the rels file
        # Python zipfile doesn't allow overwriting, so we need to rebuild
    
    # Rebuild properly with overwriting
    os.rename(tmp, tmp+'_old')
    import io
    
    with zf.ZipFile(tmp+'_old', 'r') as zin:
        with zf.ZipFile(output_path, 'w', zf.ZIP_DEFLATED) as zout:
            skip = {'word/_rels/document.xml.rels', 'word/document.xml'}
            for item in zin.infolist():
                if item.filename not in skip:
                    zout.writestr(item, zin.read(item.filename))
            
            # Add SmartArt diagram files
            for fname, content in smartart_files.items():
                zout.writestr(f'word/diagrams/{fname}', content)
            
            # Write updated rels
            try:
                rels_orig = zin.read('word/_rels/document.xml.rels').decode('utf-8')
            except:
                rels_orig = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n</Relationships>'
            
            rels_new = rels_orig.replace('</Relationships>',
                '  <Relationship Id="rId_dm" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramData" Target="diagrams/data1.xml"/>\n'
                '  <Relationship Id="rId_lo" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramLayout" Target="diagrams/layout1.xml"/>\n'
                '  <Relationship Id="rId_qs" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramQuickStyle" Target="diagrams/quickStyle1.xml"/>\n'
                '  <Relationship Id="rId_cs" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/diagramColors" Target="diagrams/colors1.xml"/>\n'
                '</Relationships>')
            zout.writestr('word/_rels/document.xml.rels', rels_new)
            
            # Modify document.xml to inject SmartArt paragraph
            doc_xml = zin.read('word/document.xml').decode('utf-8')
            # Add SmartArt before </w:body>
            doc_xml = doc_xml.replace('</w:body>', smartart_para_xml + '</w:body>')
            zout.writestr('word/document.xml', doc_xml)
    
    os.remove(tmp+'_old')


# ════════════════════════════════════════════════════════════════
# MAIN DOCUMENT BUILDER
# ════════════════════════════════════════════════════════════════

def build_document(data):
    doc = Document()
    
    # Remove default styles
    from docx.oxml.ns import nsmap
    
    # ── Page setup: landscape, small margins ──────────────────
    from docx.oxml import OxmlElement
    for section in doc.sections:
        section.page_width  = Cm(29.7)
        section.page_height = Cm(21.0)
        section.top_margin    = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin   = Cm(1.27)
        section.right_margin  = Cm(1.27)
        # Set landscape
        sectPr = section._sectPr
        pgSz = sectPr.find(qn('w:pgSz'))
        if pgSz is None:
            pgSz = OxmlElement('w:pgSz')
            sectPr.append(pgSz)
        pgSz.set(qn('w:w'), '16838')
        pgSz.set(qn('w:h'), '11906')
        pgSz.set(qn('w:orient'), 'landscape')
    
    TW_CM = 27.16  # usable width in cm (29.7 - 2*1.27)
    
    def add_table_centered(doc, cols):
        """Add table with given number of columns"""
        table = doc.add_table(rows=0, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_bidi(table)
        # Set table width to 100%
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblW')): tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)
        return table
    
    def add_centered_para(doc, text, sz=20, bold=True, color='1F3864'):
        p = doc.add_paragraph()
        set_para_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = bold
        run.font.name = FONT_NAME
        run.font.size = Pt(sz)
        run.font.color.rgb = rgb(color)
        return p
    
    def add_page_break(doc):
        p = doc.add_paragraph()
        run = p.add_run()
        from docx.oxml import OxmlElement
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)
    
    pids = [k for k, v in (data.get('products') or {}).items() if v.get('name')]
    
    # ════ TITLE PAGE ══════════════════════════════════════════
    for _ in range(8): doc.add_paragraph()
    add_centered_para(doc, 'قسم المشاريع التنموية', sz=22, bold=True, color=C['DARK'])
    add_centered_para(doc, 'نموذج طرح دراسة مشروع', sz=18, bold=True, color='000000')
    add_page_break(doc)
    
    # ════ 1. SUMMARY ══════════════════════════════════════════
    t = add_table_centered(doc, 2)
    section_hdr_row(t, 'ملخص معلومات المشروع', C['DARK'])
    rows_data = [
        ('فكرة المشروع',                       data.get('projectIdea','')),
        ('اسم مقدم المشروع',                   data.get('applicantName','')),
        ('رقم الهاتف',                          data.get('applicantPhone','')),
        ('إجمالي التكاليف التأسيسية',           data.get('summary',{}).get('foundingTotal','$0')),
        ('إجمالي الإيرادات المتوقعة (سنوياً)',  data.get('summary',{}).get('revenueAnnual','$0')),
        ('إجمالي التكاليف التشغيلية (سنوياً)',  data.get('summary',{}).get('opsAnnual','$0')),
        ('إجمالي التكاليف الثابتة (سنوياً)',    data.get('summary',{}).get('fixedAnnual','$0')),
        ('الاهتلاك (سنوياً)',                   data.get('summary',{}).get('depreciation','$0')),
        ('الربح الصافي (سنوياً)',                data.get('summary',{}).get('netProfit','$0')),
        ('عدد الموظفين في المشروع',             str(data.get('summary',{}).get('employees','0'))),
    ]
    for i,(lbl,val) in enumerate(rows_data):
        row = t.add_row()
        add_cell(row.cells[0], lbl, C['SLBL'], C['BLACK'], bold=True, align='right')
        add_cell(row.cells[1], val, C['SVAL'], C['BLACK'], bold=True)
    add_page_break(doc)
    
    # ════ 2. FOUNDING ═════════════════════════════════════════
    founding = data.get('foundingRows') or []
    if founding:
        t = add_table_centered(doc, 8)
        section_hdr_row(t, 'التكاليف التأسيسية', C['DARK'])
        col_hdr_row(t, ['#','الصنف','البيان','الاهتلاك','العدد','ملاحظات','التكلفة للواحدة','التكلفة الإجمالية'], C['HDR_BLU'], C['BLACK'])
        tot = 0
        for i, r in enumerate(founding):
            row = t.add_row()
            fill = C['WHITE'] if i%2==0 else C['REVEN']
            add_cell(row.cells[0], str(i+1),          fill, C['BLACK'], bold=True)
            add_cell(row.cells[1], r.get('cat',''),    fill, C['BLACK'], bold=True)
            add_cell(row.cells[2], r.get('bayan',''),  fill, C['BLACK'], bold=True)
            add_cell(row.cells[3], '✓' if r.get('dep') else '', fill, '70AD47' if r.get('dep') else C['BLACK'], bold=True)
            add_cell(row.cells[4], str(r.get('qty','')), fill, C['BLACK'], bold=True)
            add_cell(row.cells[5], r.get('notes',''),  fill, C['BLACK'], bold=True)
            add_cell(row.cells[6], fM(r.get('price',0)), fill, C['BLACK'], bold=True)
            add_cell(row.cells[7], fM(r.get('total',0)), fill, C['BLACK'], bold=True)
            v = float(str(r.get('total',0)).replace('$','').replace(',','') or 0)
            tot += v
        total_row(t, 'الإجمالي', fM(tot), 8, C['DARK'])
        add_page_break(doc)
    
    # ════ 3. PRODUCTS ═════════════════════════════════════════
    if pids:
        t = add_table_centered(doc, 3)
        section_hdr_row(t, 'جدول المنتجات', C['GREY'], C['BLACK'])
        col_hdr_row(t, ['البيان','الواحدة','المكونات'], C['HDR_BLK'], C['WHITE'])
        products = data.get('products') or {}
        for pi, pid in enumerate(pids):
            p = products[pid]
            pc = PROD_PAL[pi % len(PROD_PAL)]
            comps = [c for c in (p.get('components') or []) if c]
            for ci, comp in enumerate(comps):
                row = t.add_row()
                if ci == 0:
                    add_cell(row.cells[0], p.get('name',''), pc['d'], C['BLACK'], bold=True)
                    add_cell(row.cells[1], p.get('unit',''), pc['d'], C['BLACK'], bold=True)
                else:
                    add_cell(row.cells[0], '', pc['d'], C['BLACK'])
                    add_cell(row.cells[1], '', pc['d'], C['BLACK'])
                add_cell(row.cells[2], comp, pc['d'], C['BLACK'], bold=True)
        add_page_break(doc)
    
    # ════ 4. REVENUE ══════════════════════════════════════════
    if pids:
        n_cols = 2 + len(MONTHS)  # 14
        t = add_table_centered(doc, n_cols)
        section_hdr_row(t, 'الإيرادات المتوقعة', C['GREY'], C['BLACK'])
        col_hdr_row(t, ['البيان','الواحدة'] + MONTHS, C['HDR_BLK'], C['WHITE'])
        products = data.get('products') or {}
        rev_data = data.get('revenueData') or {}
        for pi, pid in enumerate(pids):
            p = products[pid]
            pc = PROD_PAL[pi % len(PROD_PAL)]
            rev = rev_data.get(pid) or [{}]*12
            if isinstance(rev, list): pass
            else: rev = [{}]*12
            # Qty row
            row = t.add_row()
            add_cell(row.cells[0], p.get('name',''), pc['d'], C['WHITE'], bold=True)
            add_cell(row.cells[1], p.get('unit',''), pc['d'], C['WHITE'], bold=True)
            for m in range(12):
                mo = rev[m] if m < len(rev) else {}
                fill = pc['o'] if m%2==0 else pc['e']
                add_cell(row.cells[m+2], fN(mo.get('qty',0)), fill, C['BLACK'], bold=True)
            # Price row
            row = t.add_row()
            add_cell(row.cells[0], 'سعر مبيع الواحدة', pc['d'], C['WHITE'], bold=True)
            add_cell(row.cells[1], '$', pc['d'], C['WHITE'], bold=True)
            for m in range(12):
                mo = rev[m] if m < len(rev) else {}
                fill = pc['o'] if m%2==0 else pc['e']
                add_cell(row.cells[m+2], fN(mo.get('unitPrice',0)), fill, C['BLACK'], bold=True)
            # Total row
            row = t.add_row()
            add_cell(row.cells[0], 'سعر المبيع الإجمالي', pc['d'], C['WHITE'], bold=True)
            add_cell(row.cells[1], '$', pc['d'], C['WHITE'], bold=True)
            for m in range(12):
                mo = rev[m] if m < len(rev) else {}
                fill = pc['o'] if m%2==0 else pc['e']
                add_cell(row.cells[m+2], fN(mo.get('total',0)), fill, C['BLACK'], bold=True)
        # Grand total
        row = t.add_row()
        # Merge first 2 cells
        row.cells[0].merge(row.cells[1])
        add_cell(row.cells[0], 'الإجمالي', C['HDR_BLK'], C['WHITE'], bold=True)
        s = data.get('summary') or {}
        # Compute month totals
        for m in range(12):
            t_val = 0
            for pid in pids:
                rev = (data.get('revenueData') or {}).get(pid) or []
                if m < len(rev):
                    mo = rev[m] if isinstance(rev,list) else {}
                    try: t_val += float(str(mo.get('total',0)).replace('$','').replace(',','') or 0)
                    except: pass
            add_cell(row.cells[m+2], fM(t_val), C['HDR_BLK'], C['WHITE'], bold=True)
        add_page_break(doc)
    
    # ════ 5. OPS ══════════════════════════════════════════════
    if pids:
        n_cols = 3 + len(MONTHS)  # 15
        t = add_table_centered(doc, n_cols)
        section_hdr_row(t, 'التكاليف التشغيلية', C['GREY'], C['BLACK'])
        col_hdr_row(t, ['البيان','الواحدة','التفاصيل'] + MONTHS, C['HDR_BLK'], C['WHITE'])
        products = data.get('products') or {}
        ops_data = data.get('opsData') or {}
        for pi, pid in enumerate(pids):
            p = products[pid]
            pc = PROD_PAL[pi % len(PROD_PAL)]
            comps = [c for c in (p.get('components') or []) if c]
            opsD = ops_data.get(pid) or ops_data.get(str(pid)) or {}
            for ci, comp in enumerate(comps):
                row = t.add_row()
                if ci == 0:
                    add_cell(row.cells[0], p.get('name',''), pc['d'], C['BLACK'], bold=True)
                    add_cell(row.cells[1], p.get('unit',''), pc['d'], C['BLACK'], bold=True)
                else:
                    add_cell(row.cells[0], '', pc['d'], C['BLACK'])
                    add_cell(row.cells[1], '', pc['d'], C['BLACK'])
                add_cell(row.cells[2], comp, pc['d'], C['BLACK'], bold=True)
                for m in range(12):
                    key = f'{ci}_{m}'
                    fill = pc['o'] if m%2==0 else pc['e']
                    add_cell(row.cells[m+3], fN(opsD.get(key,0)), fill, C['BLACK'], bold=True)
            # Subtotal row
            row = t.add_row()
            row.cells[0].merge(row.cells[2])
            add_cell(row.cells[0], f'إجمالي {p.get("name","")}', C['SUB'], C['WHITE'], bold=True, align='right')
            for m in range(12):
                key = f'sub_{m}'
                add_cell(row.cells[m+3], fM(opsD.get(key,0)), C['SUB'], C['WHITE'], bold=True)
        # Grand total
        row = t.add_row()
        row.cells[0].merge(row.cells[2])
        add_cell(row.cells[0], 'الإجمالي', C['HDR_BLK'], C['WHITE'], bold=True, align='right')
        for m in range(12):
            t_val = 0
            for pid in pids:
                opsD = (ops_data.get(pid) or ops_data.get(str(pid)) or {})
                try: t_val += float(str(opsD.get(f'sub_{m}',0)).replace('$','').replace(',','') or 0)
                except: pass
            add_cell(row.cells[m+3], fM(t_val), C['HDR_BLK'], C['WHITE'], bold=True)
        add_page_break(doc)
    
    # ════ 6. HR ═══════════════════════════════════════════════
    hr_rows = data.get('hrRows') or []
    if hr_rows:
        t = add_table_centered(doc, 7)
        section_hdr_row(t, 'الموارد البشرية', C['DARK'])
        col_hdr_row(t, ['#','المنصب','النوع','العدد','تابع لـ','الراتب الشهري الفردي','الراتب الشهري الإجمالي'], C['HDR_BLU'], C['BLACK'])
        tot = 0
        for i, r in enumerate(hr_rows):
            row = t.add_row()
            fill = C['WHITE'] if i%2==0 else C['REVEN']
            add_cell(row.cells[0], str(i+1),                  fill, C['BLACK'], bold=True)
            add_cell(row.cells[1], r.get('position',''),       fill, C['BLACK'], bold=True)
            add_cell(row.cells[2], r.get('type',''),           fill, C['BLACK'], bold=True)
            add_cell(row.cells[3], str(r.get('qty','')),       fill, C['BLACK'], bold=True)
            add_cell(row.cells[4], r.get('reports','') or '—', fill, C['BLACK'], bold=True)
            add_cell(row.cells[5], fM(r.get('salary',0)),      fill, C['BLACK'], bold=True)
            add_cell(row.cells[6], fM(r.get('total',0)),       fill, C['BLACK'], bold=True)
            try: tot += float(str(r.get('total',0)).replace('$','').replace(',','') or 0)
            except: pass
        # Total row - center "الإجمالي"
        total_row(t, 'الإجمالي', fM(tot), 7, C['DARK'])
        add_page_break(doc)
    
    # ════ 7. FIXED COSTS ══════════════════════════════════════
    fixed = data.get('fixedRows') or []
    if fixed:
        t = add_table_centered(doc, 7)
        section_hdr_row(t, 'التكاليف الثابتة', C['DARK'])
        col_hdr_row(t, ['#','الصنف','البيان','العدد','ملاحظات','التكلفة الشهرية للواحدة','التكلفة الشهرية الإجمالية'], C['HDR_BLU'], C['BLACK'])
        tot = 0
        for i, r in enumerate(fixed):
            row = t.add_row()
            fill = C['WHITE'] if i%2==0 else C['REVEN']
            add_cell(row.cells[0], str(i+1),             fill, C['BLACK'], bold=True)
            add_cell(row.cells[1], r.get('cat',''),       fill, C['BLACK'], bold=True)
            add_cell(row.cells[2], r.get('bayan',''),     fill, C['BLACK'], bold=True)
            add_cell(row.cells[3], str(r.get('qty','')), fill, C['BLACK'], bold=True)
            add_cell(row.cells[4], r.get('notes',''),     fill, C['BLACK'], bold=True)
            add_cell(row.cells[5], fM(r.get('price',0)), fill, C['BLACK'], bold=True)
            add_cell(row.cells[6], fM(r.get('total',0)), fill, C['BLACK'], bold=True)
            try: tot += float(str(r.get('total',0)).replace('$','').replace(',','') or 0)
            except: pass
        total_row(t, 'الإجمالي', fM(tot), 7, C['DARK'])
        add_page_break(doc)
    
    # ════ 8. DEPRECIATION ════════════════════════════════════
    dep_rows = data.get('depRows') or []
    if dep_rows:
        t = add_table_centered(doc, 8)
        section_hdr_row(t, 'الاهتلاك', C['DARK'])
        col_hdr_row(t, ['#','الصنف','البيان','نسبة الاهتلاك','العدد','ملاحظات','قيمة الاهتلاك للواحدة','قيمة الاهتلاك الإجمالية'], C['HDR_BLU'], C['BLACK'])
        tot = 0
        for i, r in enumerate(dep_rows):
            row = t.add_row()
            fill = C['WHITE'] if i%2==0 else C['REVEN']
            add_cell(row.cells[0], str(i+1),                         fill, C['BLACK'], bold=True)
            add_cell(row.cells[1], r.get('cat',''),                   fill, C['BLACK'], bold=True)
            add_cell(row.cells[2], r.get('bayan',''),                 fill, C['BLACK'], bold=True)
            add_cell(row.cells[3], str(r.get('pct','0')) + ' %',      fill, C['BLACK'], bold=True)
            add_cell(row.cells[4], str(r.get('qty','')),              fill, C['BLACK'], bold=True)
            add_cell(row.cells[5], r.get('notes',''),                 fill, C['BLACK'], bold=True)
            add_cell(row.cells[6], fM(r.get('perUnit',0)),            fill, C['BLACK'], bold=True)
            add_cell(row.cells[7], fM(r.get('total',0)),              fill, C['BLACK'], bold=True)
            try: tot += float(str(r.get('total',0)).replace('$','').replace(',','') or 0)
            except: pass
        total_row(t, 'إجمالي قيمة الاهتلاك', fM(tot), 8, C['DARK'])
    
    return doc


if __name__ == '__main__':
    import docx.oxml.shared
    from docx.oxml import register_element_cls
    # monkey-patch for break
    import docx.text.run as _run_mod
    
    data_file = sys.argv[1]
    out_file  = sys.argv[2]
    
    with open(data_file) as f:
        data = json.load(f)
    
    doc = build_document(data)
    
    hr_rows = data.get('hrRows') or []
    if hr_rows:
        embed_smartart_in_docx(doc, hr_rows, out_file)
    else:
        doc.save(out_file)
    
    print(f'OK:{out_file}')
